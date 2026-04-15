import streamlit as st
from datetime import datetime
from docx import Document

st.set_page_config(page_title="Gerador de Comprovante", layout="centered")

st.title("📄 Gerador de Comprovante PIX")

# memória da chave pix
if "chave_pix" not in st.session_state:
    st.session_state.chave_pix = ""

# upload da imagem
imagem = st.file_uploader("📷 Envimport streamlit as st
from datetime import datetime
from docx import Document
import os

st.set_page_config(page_title="Gerador de Comprovante", layout="centered")

st.title("📄 Gerador de Comprovante PIX")

# memória da chave pix
if "chave_pix" not in st.session_state:
    st.session_state.chave_pix = ""

# upload da imagem
imagem = st.file_uploader("📷 Envie a imagem do comprovante", type=["png", "jpg", "jpeg"])

# inputs do pagador
pagador_nome = st.text_input("Nome do pagador")
pagador_cpf = st.text_input("CPF do pagador")
valor = st.text_input("Valor (ex: 1290,00)")

# chave pix com memória
chave_pix = st.text_input("Chave Pix", value=st.session_state.chave_pix)

if chave_pix:
    st.session_state.chave_pix = chave_pix

# DADOS SIMULADOS (depois podemos colocar OCR real)
recebedor_nome = "Recebedor da Imagem"
recebedor_doc = "***.000.000-**"
banco = "MERCADO PAGO IP LTDA."

# botão gerar documento
if st.button("🚀 Gerar Documento"):

    if not pagador_nome or not pagador_cpf or not valor:
        st.warning("Preencha todos os campos")

    else:
        # verifica se o modelo existe
        if not os.path.exists("modelo.docx"):
            st.error("❌ Arquivo modelo.docx não encontrado no repositório!")
            st.stop()

        # abre o modelo
        doc = Document("modelo.docx")

        agora = datetime.now()
        data_formatada = agora.strftime("%d/%m/%Y - %H:%M:%S")
        data_transacao = agora.strftime("%d/%m/%Y %H:%M:%S")

        # substituição dos campos
        for p in doc.paragraphs:
            p.text = p.text.replace("{{pagador_nome}}", pagador_nome)
            p.text = p.text.replace("{{pagador_cpf}}", pagador_cpf)
            p.text = p.text.replace("{{valor}}", valor)
            p.text = p.text.replace("{{chave_pix}}", chave_pix)
            p.text = p.text.replace("{{recebedor_nome}}", recebedor_nome)
            p.text = p.text.replace("{{recebedor_doc}}", recebedor_doc)
            p.text = p.text.replace("{{banco}}", banco)
            p.text = p.text.replace("{{data_hora}}", data_formatada)
            p.text = p.text.replace("{{data_transacao}}", data_transacao)

        # salva o arquivo final
        arquivo_docx = "comprovante.docx"
        doc.save(arquivo_docx)

        # botão de download
        with open(arquivo_docx, "rb") as f:
            st.success("✅ Documento gerado com sucesso!")
            st.download_button(
                label="⬇️ Baixar Documento",
                data=f,
                file_name="comprovante.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )ie a imagem do comprovante", type=["png", "jpg", "jpeg"])

# inputs do pagador
pagador_nome = st.text_input("Nome do pagador")
pagador_cpf = st.text_input("CPF do pagador")
valor = st.text_input("Valor (ex: 1290,00)")

# chave pix com memória
chave_pix = st.text_input("Chave Pix", value=st.session_state.chave_pix)

if chave_pix:
    st.session_state.chave_pix = chave_pix

# DADOS SIMULADOS (depois colocamos OCR real)
recebedor_nome = "Recebedor da Imagem"
recebedor_doc = "*.000.000-**"
banco = "MERCADO PAGO IP LTDA."

# botão gerar
if st.button("🚀 Gerar Documento"):
    if not pagador_nome or not pagador_cpf or not valor:
        st.warning("Preencha todos os campos")
    else:
import os

if not os.path.exists("modelo.docx"):
    st.error("Arquivo modelo.docx não encontrado no repositório!")
    st.stop()

doc = Document("modelo.docx")
        agora = datetime.now()
        data_formatada = agora.strftime("%d/%m/%Y - %H:%M:%S")
        data_transacao = agora.strftime("%d/%m/%Y %H:%M:%S")

        # substituição dos campos
        for p in doc.paragraphs:
            p.text = p.text.replace("{{pagador_nome}}", pagador_nome)
            p.text = p.text.replace("{{pagador_cpf}}", pagador_cpf)
            p.text = p.text.replace("{{valor}}", valor)
            p.text = p.text.replace("{{chave_pix}}", chave_pix)
            p.text = p.text.replace("{{recebedor_nome}}", recebedor_nome)
            p.text = p.text.replace("{{recebedor_doc}}", recebedor_doc)
            p.text = p.text.replace("{{banco}}", banco)
            p.text = p.text.replace("{{data_hora}}", data_formatada)
            p.text = p.text.replace("{{data_transacao}}", data_transacao)

        arquivo_docx = "comprovante.docx"
        doc.save(arquivo_docx)

        # botão download
        with open(arquivo_docx, "rb") as f:
            st.success("✅ Documento gerado com sucesso!")
            st.download_button(
                label="⬇️ Baixar Documento",
                data=f,
                file_name="comprovante.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
